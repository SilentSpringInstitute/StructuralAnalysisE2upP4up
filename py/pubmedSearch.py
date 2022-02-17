import toolbox
import PubMed
import pandas as pd
import xlsxwriter
from docx import Document
from os import path

class pubmedSearch:
    def __init__(self, p_tab_chem, l_term_search, pr_out, email):
        
        self.pr_out = pr_out
        self.l_terms = l_term_search
        self.p_tab_chem = p_tab_chem
        self.email = email
        self.l_confidence = ["high", "medium", "low"]
        

    def load_tab_chem(self):
        
        # need to check the sheet name
        xl = pd.ExcelFile(self.p_tab_chem)
        l_sheet = xl.sheet_names
        if "ToxPrint_QSAR" in l_sheet:
            d_chem = toolbox.loadExcelSheet(self.p_tab_chem, name_sheet='ToxPrint_QSAR', k_head = "CASRN")
        elif "Sheet1" in l_sheet:
            d_chem = toolbox.loadExcelSheet(self.p_tab_chem, name_sheet='Sheet1', k_head = "CASRN")
        else:
            return 
        
        self.d_chem = d_chem

        # check if query is in the dictionnary
        for chem in self.d_chem.keys():
            if not "PubMed search chemical" in list(self.d_chem[chem].keys()):
                query = "(\"%s\" OR \"%s\")"%(chem, self.d_chem[chem]["Chemical name"])
                self.d_chem[chem]["PubMed search chemical"] = query
            else:
                if pd.isna(self.d_chem[chem]["PubMed search chemical"]) == True:
                    query = "(\"%s\" OR \"%s\")"%(chem, self.d_chem[chem]["Chemical name"])
                    self.d_chem[chem]["PubMed search chemical"] = query
                    

    def do_search_bychem(self):
        
        p_xlx = "%schem_search.xlsx"%(self.pr_out)
        if path.exists(p_xlx):
            return 
        workbook = xlsxwriter.Workbook(p_xlx)
        worksheet = workbook.add_worksheet()
        l_headers = ["CASRN", "Chemical name", "Confidence", "Query", "Nb Result", "Article"]
        l_chem = list(self.d_chem.keys())
        # write header in the xlxs            
        [worksheet.write(0, i_col, l_headers[i_col]) for i_col in range(0, 4)]
            
        row = 1
        imax = len(list(self.d_chem.keys()))
        i = 0
        while i < imax:
            # check first if query is in the directory
            if pd.isna(self.d_chem[l_chem[i]]["PubMed search chemical"]) == True:
                i = i + 1
                continue
            
            # check confidence
            if "Confidence\nhigh (perf > 0.5, AD > 0.75, nb signif >= 3)\nLow ( AD < 0.75 and nb signif < 3)\nMedium (AD < 0.75 or signif < 3 )" in list(self.d_chem[l_chem[i]].keys()):
                if not self.d_chem[l_chem[i]]["Confidence\nhigh (perf > 0.5, AD > 0.75, nb signif >= 3)\nLow ( AD < 0.75 and nb signif < 3)\nMedium (AD < 0.75 or signif < 3 )"] in self.l_confidence:
                    i = i + 1    
                    continue
            
            c_query = PubMed.searching(self.email, "")
            c_query.buildNewQuery(l_chem[i], self.d_chem[l_chem[i]]["PubMed search chemical"])
            c_query.search(w=False)
            
            # check process
            print(l_chem[i], c_query.count, c_query.err)
                
            if c_query.err == 0:
                # add casrn
                worksheet.write(row, 0, l_chem[i])
                # add name
                worksheet.write(row, 1, self.d_chem[l_chem[i]]["Chemical name"])
                # confidence
                if "Confidence\nhigh (perf > 0.5, AD > 0.75, nb signif >= 3)\nLow ( AD < 0.75 and nb signif < 3)\nMedium (AD < 0.75 or signif < 3 )" in list(self.d_chem[l_chem[i]].keys()):
                    worksheet.write(row, 2, self.d_chem[l_chem[i]]["Confidence\nhigh (perf > 0.5, AD > 0.75, nb signif >= 3)\nLow ( AD < 0.75 and nb signif < 3)\nMedium (AD < 0.75 or signif < 3 )"])
                else:
                    worksheet.write(row, 2, "NA")
                # add query 
                worksheet.write(row, 3, c_query.Query)
                # add nb results
                worksheet.write(row, 4, c_query.count)
                # add list of article
                worksheet.write(row, 5, "\n".join(c_query.l_titles))
            else:
                # add casrn
                worksheet.write(row, 0, l_chem[i])
                # add name
                worksheet.write(row, 1, self.d_chem[l_chem[i]]["Chemical name"])
                # confidence
                if "Confidence\nhigh (perf > 0.5, AD > 0.75, nb signif >= 3)\nLow ( AD < 0.75 and nb signif < 3)\nMedium (AD < 0.75 or signif < 3 )" in list(self.d_chem[l_chem[i]].keys()):
                    worksheet.write(row, 2, self.d_chem[l_chem[i]]["Confidence\nhigh (perf > 0.5, AD > 0.75, nb signif >= 3)\nLow ( AD < 0.75 and nb signif < 3)\nMedium (AD < 0.75 or signif < 3 )"])
                else:
                    worksheet.write(row, 2, "NA")
                # add query 
                worksheet.write(row, 3, c_query.Query)
                # add nb results
                worksheet.write(row, 4, "NA")
                # add list of article
                worksheet.write(row, 5, "Error request")
            row = row + 1
            i = i + 1
        workbook.close()
 
    def do_search_combine_terms(self, write_abstract=False):
        
        ## write in excel to have several lines in article list
        p_xlx = "%schem_combined.xlsx"%(self.pr_out)
        #if path.exists(p_xlx):
        #    return
        
        if write_abstract == True:
            #open words docx
            p_docxs = "%schem_combined.docx"%(self.pr_out)
            p_docx_high = "%schem_high_combined.docx"%(self.pr_out)
            document = Document()
            document_high_pred = Document()
        
        workbook = xlsxwriter.Workbook(p_xlx)
        worksheet = workbook.add_worksheet()
        l_headers = ["CASRN", "Chemical name", "Pred", "AD", "Confidence", "Query", "Nb Result", "Article"]
        # wrtie header in the xlxs            
        [worksheet.write(0, i_col, l_headers[i_col]) for i_col in range(0, len(l_headers))]
        row = 1
        
        l_chem = list(self.d_chem.keys())
        i = 0
        imax = len(list(self.d_chem.keys())) 
        while i < imax:
            l_queries = []
            if pd.isna(self.d_chem[l_chem[i]]["PubMed search chemical"]) == True:
                i = i + 1
                continue
                
            for term in self.l_terms: 
                c_Q_builder = PubMed.searching(self.email, "")
                c_Q_builder.buildNewQuery(l_chem[i], self.d_chem[l_chem[i]]["PubMed search chemical"])
                c_Q_builder.buildingQuery(term, "AND")
                l_queries.append(c_Q_builder.Query)
            
            # here I do the search
            c_search = PubMed.searching(self.email, "")
            c_search.buildNewQuery(l_chem[i], l_queries[0])
            i_q = 1
            i_q_max = len(l_queries)
            while i_q < i_q_max:
                c_search.buildingQuery(l_queries[i_q], "OR")
                i_q = i_q + 1
            c_search.search(w=False)
            
            if c_search.err == 0:
                # add casrn
                worksheet.write(row, 0, l_chem[i])
                # add name
                worksheet.write(row, 1, self.d_chem[l_chem[i]]["Chemical name"])
                # confidence
                if "Confidence\nhigh (perf > 0.5, AD > 0.75, nb signif >= 3)\nLow ( AD < 0.75 and nb signif < 3)\nMedium (AD < 0.75 or signif < 3 )" in list(self.d_chem[l_chem[i]].keys()):
                    worksheet.write(row, 4, self.d_chem[l_chem[i]]["Confidence\nhigh (perf > 0.5, AD > 0.75, nb signif >= 3)\nLow ( AD < 0.75 and nb signif < 3)\nMedium (AD < 0.75 or signif < 3 )"])
                    confidence = self.d_chem[l_chem[i]]["Confidence\nhigh (perf > 0.5, AD > 0.75, nb signif >= 3)\nLow ( AD < 0.75 and nb signif < 3)\nMedium (AD < 0.75 or signif < 3 )"]    
                else:
                    worksheet.write(row, 4, "NA")
                    confidence = "NA"
                
                # add pred
                worksheet.write(row, 2, self.d_chem[l_chem[i]]["Pred RF balanced"])
                # add AD
                worksheet.write(row, 3, self.d_chem[l_chem[i]]["AD distance to first neighbord"])
                # add query 
                worksheet.write(row, 5, c_search.Query)
                # add nb results
                worksheet.write(row, 6, c_search.count)
                # add list of article
                worksheet.write(row, 7, "\n".join(c_search.l_titles))
                
                ## HERE write the docx
                if write_abstract == True:
                    i_a = 0
                    i_a_max = len(c_search.l_titles)
                    document.add_heading("Chemical: %s\n%s\nPrediction Score: %s\nConfidence: %s\n"%(l_chem[i], self.d_chem[l_chem[i]]["Chemical name"], self.d_chem[l_chem[i]]["Pred RF balanced"], confidence), level=1)
                    p = document.add_paragraph()
                    p.add_run("Number of abstract: %s\n"%(c_search.count)).bold=True
                    
                    if confidence == "high":
                        document_high_pred.add_heading("Chemical: %s\n%s\nPrediction Score: %s\nConfidence: %s\n"%(l_chem[i], self.d_chem[l_chem[i]]["Chemical name"], self.d_chem[l_chem[i]]["Pred RF balanced"], confidence), level=1)
                        p_high = document_high_pred.add_paragraph()
                        p_high.add_run("Number of abstract: %s\n"%(c_search.count)).bold=True
                    
                    while i_a < i_a_max:
                        p.add_run("%s. Title: "%(i_a + 1)).bold=True
                        p.add_run(c_search.l_titles[i_a] + "\n")
                        p.add_run("Abstract: ").bold=True
                        p.add_run(c_search.l_abstracts[i_a] + "\n\n")
                        
                        #write file with only high confidence cheminal
                        if confidence == "high":
                            p_high.add_run("%s. Title: "%(i_a + 1)).bold=True
                            p_high.add_run(c_search.l_titles[i_a] + "\n")
                            p_high.add_run("Abstract: ").bold=True
                            p_high.add_run(c_search.l_abstracts[i_a] + "\n\n")
                        
                        i_a = i_a + 1
            else:
                # add casrn
                worksheet.write(row, 0, l_chem[i])
                # add name
                worksheet.write(row, 1, self.d_chem[l_chem[i]]["Chemical name"])
                # confidence
                if "Confidence\nhigh (perf > 0.5, AD > 0.75, nb signif >= 3)\nLow ( AD < 0.75 and nb signif < 3)\nMedium (AD < 0.75 or signif < 3 )" in list(self.d_chem[l_chem[i]].keys()):
                    worksheet.write(row, 4, self.d_chem[l_chem[i]]["Confidence\nhigh (perf > 0.5, AD > 0.75, nb signif >= 3)\nLow ( AD < 0.75 and nb signif < 3)\nMedium (AD < 0.75 or signif < 3 )"])
                else:
                    worksheet.write(row, 4, "NA")
                
                # add pred
                worksheet.write(row, 2, self.d_chem[l_chem[i]]["Pred RF balanced"])
                # add AD
                worksheet.write(row, 3, self.d_chem[l_chem[i]]["AD distance to first neighbord"])
                # add query 
                worksheet.write(row, 5, c_search.Query)
                # add nb results
                worksheet.write(row, 6, "NA")
                # add list of article
                worksheet.write(row, 7, "Error request")
                
            row = row + 1
            i = i + 1
        
        if write_abstract == True: 
            document.save(p_docxs)
            document_high_pred.save(p_docx_high)
        workbook.close()
        
    def do_search_by_term(self):
        
        ## write in excel to have several lines in article list
        for term in self.l_terms:
            p_xlx = "%schem_%s.xlsx"%(self.pr_out, term)
            if path.exists(p_xlx):
                continue
            workbook = xlsxwriter.Workbook(p_xlx)
            worksheet = workbook.add_worksheet()
            l_headers = ["CASRN", "Query", "Nb Result", "Article"]
            l_chem = list(self.d_chem.keys())
            # wrtie header in the xlxs            
            [worksheet.write(0, i_col, l_headers[i_col]) for i_col in range(0, 4)]
            
            row = 1
            imax = len(list(self.d_chem.keys()))
            i = 0
            while i < imax:
                if pd.isna(self.d_chem[l_chem[i]]["PubMed search chemical"]) == True:
                    i = i + 1
                    continue
                
                c_query = PubMed.searching(self.email, "")
                c_query.buildNewQuery(l_chem[i], self.d_chem[l_chem[i]]["PubMed search chemical"])
                c_query.buildingQuery(term, "AND")
                c_query.search(w=False)
                
                # add casrn
                worksheet.write(row, 0, l_chem[i])
                # add query 
                worksheet.write(row, 1, c_query.Query)
                # add nb results
                worksheet.write(row, 2, c_query.count)
                # add list of article
                worksheet.write(row, 3, "\n".join(c_query.l_titles))
                
                row = row + 1
                i = i + 1
            
            workbook.close()

    def do_search(self, review = "on", date = "", write_abstract = True):
        self.load_tab_chem()
        #self.do_search_bychem()
        #self.do_search_by_term()
        self.do_search_combine_terms(write_abstract=write_abstract)
                


